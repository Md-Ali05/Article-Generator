from langchain.llms import CTransformers
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
import streamlit as st 
import os
from docx import Document
from docx.shared import Inches
import io

def load_llm(max_tokens, prompt_template):
    # Load the locally downloaded model here
    llm = CTransformers(
        model = "llama-2-7b-chat.ggmlv3.q8_0.bin",
        model_type="llama",
        max_new_tokens = max_tokens,
        temperature = 0.65
        
    )
    
    llm_chain = LLMChain(
        llm=llm,
        prompt=PromptTemplate.from_template(prompt_template)
    )
    print(llm_chain)
    return llm_chain


def create_word_docx(user_input, paragraph):
    # Create a new Word document
    doc = Document()

    # Add the user input to the document
    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)
    return doc

st.set_page_config(layout="wide")
def main():

       st.title("Article Generator App using Meta's Llama 2 open source LLM model")

       user_input = st.text_input("Please enter the topic for the article you want to generate!. Note: The generated article will have around 100-150 words")
       

       if len(user_input) > 0:

        col1, col2= st.columns([3,1])

        with col1:
            st.subheader("You can repeatedly prompt for article generation for a better response")
            st.write("Topic of the article is: " + user_input)
            
            prompt_template = """Write me a 150 word blog post about {user_input}. Make it formal and concise."""
            llm_call = load_llm(max_tokens=4096, prompt_template=prompt_template)
            print(llm_call)
            result = llm_call(user_input)
            if len(result) > 0:
                st.info("Your article has been been generated successfully!")
                st.write(result)
            else:
                st.error("Your article couldn't be generated!")

        with col2:
            st.subheader("Final Article to Download")
            
            doc = create_word_docx(user_input, result['text'])

            # Save the Word document to a BytesIO buffer
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Prepare the download link
            st.download_button(
                label='Download Word Document',
                data=doc_buffer,
                file_name='document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )


if __name__ == "__main__":
    main()